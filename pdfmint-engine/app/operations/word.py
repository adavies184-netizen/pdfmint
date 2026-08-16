from __future__ import annotations

import logging
import os
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Iterable

import fitz
from lxml import etree
from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from ..settings import JOB_TIMEOUT_SECONDS

logger = logging.getLogger("pdfmint.word")
VML_NAMESPACE = "urn:schemas-microsoft-com:vml"


def _anchor_picture_to_page(inline_shape) -> None:
    """Turn python-docx's inline picture into an exact page-positioned anchor."""
    inline = inline_shape._inline
    inline.tag = qn("wp:anchor")
    inline.set("distT", "0")
    inline.set("distB", "0")
    inline.set("distL", "0")
    inline.set("distR", "0")
    inline.set("simplePos", "0")
    inline.set("relativeHeight", "0")
    inline.set("behindDoc", "1")
    inline.set("locked", "1")
    inline.set("layoutInCell", "1")
    inline.set("allowOverlap", "1")

    simple_position = OxmlElement("wp:simplePos")
    simple_position.set("x", "0")
    simple_position.set("y", "0")

    horizontal = OxmlElement("wp:positionH")
    horizontal.set("relativeFrom", "page")
    horizontal_offset = OxmlElement("wp:posOffset")
    horizontal_offset.text = "0"
    horizontal.append(horizontal_offset)

    vertical = OxmlElement("wp:positionV")
    vertical.set("relativeFrom", "page")
    vertical_offset = OxmlElement("wp:posOffset")
    vertical_offset.text = "0"
    vertical.append(vertical_offset)

    wrap_none = OxmlElement("wp:wrapNone")
    inline.insert(0, simple_position)
    inline.insert(1, horizontal)
    inline.insert(2, vertical)
    effect_extent_index = next(
        (index for index, child in enumerate(inline) if child.tag == qn("wp:effectExtent")),
        3,
    )
    inline.insert(effect_extent_index + 1, wrap_none)


def _word_font_name(pdf_font: str) -> str:
    name = str(pdf_font or "Arial").split("+")[-1]
    lowered = name.lower()
    if "helvetica" in lowered or "arial" in lowered:
        return "Arial"
    if "times" in lowered or "serif" in lowered:
        return "Times New Roman"
    if "courier" in lowered or "mono" in lowered:
        return "Courier New"
    return name.replace("-Bold", "").replace("-Italic", "") or "Arial"


def _colour_hex(value: int) -> str:
    return f"{int(value or 0) & 0xFFFFFF:06X}"


def _append_editable_text_box(paragraph, span: dict, shape_id: int) -> None:
    text = str(span.get("text") or "")
    if not text:
        return

    x0, y0, x1, y1 = [float(value) for value in span["bbox"]]
    size = max(1.0, float(span.get("size") or 11.0))
    flags = int(span.get("flags") or 0)
    font_name = _word_font_name(span.get("font", "Arial"))
    width = max(2.0, x1 - x0 + max(2.0, size * .22))
    height = max(size * 1.35, y1 - y0 + 2.0)

    run = paragraph.add_run()
    pict = OxmlElement("w:pict")
    shape = etree.Element(f"{{{VML_NAMESPACE}}}shape")
    shape.set("id", f"pdfmint-text-{shape_id}")
    shape.set("type", "#_x0000_t202")
    shape.set(
        "style",
        ";".join(
            [
                "position:absolute",
                f"margin-left:{x0:.3f}pt",
                f"margin-top:{max(0.0, y0 - size * .08):.3f}pt",
                f"width:{width:.3f}pt",
                f"height:{height:.3f}pt",
                "z-index:251659264",
                "mso-wrap-style:none",
                "mso-position-horizontal-relative:page",
                "mso-position-vertical-relative:page",
            ]
        ),
    )
    shape.set("stroked", "f")
    shape.set("filled", "f")

    textbox = etree.Element(f"{{{VML_NAMESPACE}}}textbox")
    textbox.set("inset", "0,0,0,0")
    content = OxmlElement("w:txbxContent")
    text_paragraph = OxmlElement("w:p")
    paragraph_properties = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), str(max(20, round(size * 20 * 1.15))))
    spacing.set(qn("w:lineRule"), "exact")
    paragraph_properties.append(spacing)
    text_paragraph.append(paragraph_properties)

    text_run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), font_name)
    run_properties.append(fonts)
    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), _colour_hex(span.get("color", 0)))
    run_properties.append(colour)
    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), str(max(2, round(size * 2))))
    run_properties.append(font_size)
    complex_font_size = OxmlElement("w:szCs")
    complex_font_size.set(qn("w:val"), str(max(2, round(size * 2))))
    run_properties.append(complex_font_size)
    if flags & 16:
        run_properties.append(OxmlElement("w:b"))
    if flags & 2:
        run_properties.append(OxmlElement("w:i"))
    text_run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_node.text = text
    text_run.append(text_node)
    text_paragraph.append(text_run)
    content.append(text_paragraph)
    textbox.append(content)
    shape.append(textbox)
    pict.append(shape)
    run._r.append(pict)


def _page_text_spans(page) -> list[dict]:
    spans: list[dict] = []
    for block in page.get_text("dict", sort=True).get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                if str(span.get("text") or "").strip():
                    spans.append(span)
    return spans


def _render_graphics_background(source, page_index: int, spans: list[dict], image_path: Path) -> None:
    background = fitz.open()
    try:
        background.insert_pdf(source, from_page=page_index, to_page=page_index)
        page = background[0]
        for span in spans:
            rectangle = fitz.Rect(span["bbox"])
            rectangle.x0 -= .35
            rectangle.y0 -= .35
            rectangle.x1 += .35
            rectangle.y1 += .35
            page.add_redact_annot(rectangle, fill=None, cross_out=False)
        if spans:
            page.apply_redactions(images=0, graphics=0, text=0)
        pixmap = page.get_pixmap(matrix=fitz.Matrix(2.5, 2.5), alpha=False)
        pixmap.save(str(image_path))
    finally:
        background.close()


def pdf_to_docx(pdf_path: Path, output_path: Path) -> Path:
    """Create a visually faithful Word file whose detected text stays editable."""
    source = fitz.open(str(pdf_path))
    if source.page_count < 1:
        source.close()
        raise RuntimeError("The PDF does not contain any pages.")

    document = Document()

    with tempfile.TemporaryDirectory(prefix="pdfmint-word-pages-") as temp_dir:
        try:
            for page_index, page in enumerate(source):
                if page_index == 0:
                    section = document.sections[0]
                    paragraph = document.add_paragraph()
                else:
                    section = document.add_section(WD_SECTION.NEW_PAGE)
                    paragraph = document.paragraphs[-1]

                page_width = float(page.rect.width)
                page_height = float(page.rect.height)
                section.page_width = Pt(page_width)
                section.page_height = Pt(page_height)
                section.top_margin = Pt(0)
                section.bottom_margin = Pt(0)
                section.left_margin = Pt(0)
                section.right_margin = Pt(0)
                section.header_distance = Pt(0)
                section.footer_distance = Pt(0)

                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)

                spans = _page_text_spans(page)
                image_path = Path(temp_dir) / f"page-{page_index + 1}.png"
                _render_graphics_background(source, page_index, spans, image_path)
                picture = paragraph.add_run().add_picture(
                    str(image_path),
                    width=Pt(page_width),
                    height=Pt(page_height),
                )
                _anchor_picture_to_page(picture)
                for span in spans:
                    _append_editable_text_box(paragraph, span, shape_id=(page_index + 1) * 10000 + len(paragraph.runs))

            document.save(str(output_path))
        finally:
            source.close()

    if not output_path.exists() or output_path.stat().st_size == 0:
        raise RuntimeError("DOCX conversion did not produce a valid file.")

    logger.info(
        "Hybrid editable DOCX created path=%s size_bytes=%s",
        output_path,
        output_path.stat().st_size,
    )
    return output_path


def _remove_old_doc_output(doc_path: Path) -> None:
    try:
        if doc_path.exists():
            doc_path.unlink()
    except OSError:
        logger.warning("Could not remove previous DOC output: %s", doc_path)


def _run_libreoffice_export(
    docx_path: Path,
    output_dir: Path,
    libreoffice_home: Path,
    convert_to: str,
) -> tuple[subprocess.CompletedProcess[str], Path]:
    doc_path = output_dir / f"{docx_path.stem}.doc"
    _remove_old_doc_output(doc_path)

    command = [
        "soffice",
        "--headless",
        "--invisible",
        "--nologo",
        "--nodefault",
        "--nolockcheck",
        "--nofirststartwizard",
        "--norestore",
        f"-env:UserInstallation=file://{libreoffice_home}",
        "--convert-to",
        convert_to,
        "--outdir",
        str(output_dir),
        str(docx_path),
    ]

    logger.info("LibreOffice attempt convert_to=%s", convert_to)
    logger.info("Command: %s", " ".join(command))

    result = subprocess.run(
        command,
        capture_output=True,
        text=True,
        timeout=JOB_TIMEOUT_SECONDS,
        check=False,
        env={
            **os.environ,
            "HOME": str(libreoffice_home),
            "SAL_USE_VCLPLUGIN": "svp",
            "JAVA_TOOL_OPTIONS": "-Xms16m -Xmx64m",
        },
    )

    logger.info("LibreOffice exit code: %s", result.returncode)
    logger.info("LibreOffice stdout: %s", (result.stdout or "").strip()[:4000])
    logger.info("LibreOffice stderr: %s", (result.stderr or "").strip()[:4000])
    logger.info("Output exists: %s", doc_path.exists())
    logger.info("Output size bytes: %s", doc_path.stat().st_size if doc_path.exists() else -1)

    return result, doc_path


def _candidate_doc_filters() -> Iterable[str]:
    yield "doc"
    yield 'doc:"MS Word 97"'
    yield 'doc:"MS Word 95"'
    yield 'doc:"Office Open XML Text"'


def docx_to_doc(docx_path: Path, output_dir: Path) -> Path:
    libreoffice_home = output_dir / "libreoffice-home"
    libreoffice_home.mkdir(parents=True, exist_ok=True)

    logger.info("========== DOC CONVERSION START ==========")
    logger.info("Input DOCX: %s", docx_path)
    logger.info("Input size bytes: %s", docx_path.stat().st_size if docx_path.exists() else -1)
    logger.info("Output directory: %s", output_dir)

    failures: list[str] = []

    for convert_to in _candidate_doc_filters():
        attempt_slug = (
            convert_to.replace('"', "")
            .replace(":", "-")
            .replace(" ", "-")
            .lower()
        )
        attempt_home = libreoffice_home / attempt_slug
        if attempt_home.exists():
            shutil.rmtree(attempt_home, ignore_errors=True)
        attempt_home.mkdir(parents=True, exist_ok=True)

        try:
            result, doc_path = _run_libreoffice_export(
                docx_path=docx_path,
                output_dir=output_dir,
                libreoffice_home=attempt_home,
                convert_to=convert_to,
            )
        except subprocess.TimeoutExpired:
            logger.exception("LibreOffice timed out for convert_to=%s", convert_to)
            failures.append(f"{convert_to}: timed out")
            continue

        if doc_path.exists() and doc_path.stat().st_size > 0 and result.returncode == 0:
            logger.info("DOC export succeeded with convert_to=%s", convert_to)
            logger.info("========== DOC CONVERSION END ==========")
            return doc_path

        details = (result.stderr or result.stdout or "No DOC output was produced.").strip()
        failures.append(f"{convert_to}: exit={result.returncode}; {details[:600]}")

    logger.info("========== DOC CONVERSION END ==========")
    raise RuntimeError(
        "LibreOffice could not export DOC using any supported filter. "
        + " | ".join(failures)
    )


def pdf_to_doc(pdf_path: Path, output_dir: Path, base_name: str) -> Path:
    docx_path = output_dir / f"{base_name}.docx"
    pdf_to_docx(pdf_path, docx_path)
    return docx_to_doc(docx_path, output_dir)
